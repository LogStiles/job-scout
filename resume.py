"""Distill a .docx resume into a short candidate profile for the scorer.

Extracts text from a Word resume and uses an LLM to condense it into the
same bullet-point profile format that scorer.py scores against.
"""

from __future__ import annotations

import argparse
import sys
from typing import List

import anthropic
from docx import Document as _DocxDocument
from pydantic import BaseModel, Field

from scorer import MODEL, PROFILE_PATH, Profile, save_default_profile


class DistilledProfile(BaseModel):
    """The structured profile fields distilled from a resume."""

    title_target: str = Field(
        description="The role/title the candidate is targeting, inferred from "
        "their most recent and senior experience (e.g. 'Senior Software Engineer')."
    )
    core_skills: List[str] = Field(
        description="The candidate's strongest, most relevant technical skills, "
        "backed by substantial professional (on-the-job) experience."
    )
    soft_skills: List[str] = Field(
        default_factory=list,
        description="Skills the candidate has demonstrated but WITHOUT strong "
        "professional experience — shown instead through personal/academic "
        "projects, coursework, certifications, or education. Do not repeat "
        "anything already listed in core_skills. Empty if none are evident.",
    )
    years_experience: int = Field(
        ge=0,
        description="Total years of professional experience, estimated from work history.",
    )
    location: str = Field(
        description="Where the candidate is based and any stated work-arrangement "
        "preferences (remote/hybrid/onsite). Use 'Not specified' if absent."
    )
    preferred_domains: List[str] = Field(
        default_factory=list,
        description="Industries/domains the candidate has worked in or targets "
        "(e.g. fintech, healthtech). Empty if none are evident.",
    )
    hard_disqualifiers: str = Field(
        default="None specified",
        description="Constraints that should rule a posting out, if stated on the "
        "resume (e.g. location limits, unwanted stacks). 'None specified' if absent.",
    )

    def to_profile(self) -> Profile:
        """Render into the bullet-point format scorer.build_prompt expects."""
        skills = ", ".join(self.core_skills) if self.core_skills else "Not specified"
        soft = ", ".join(self.soft_skills) if self.soft_skills else "Not specified"
        domains = (
            ", ".join(self.preferred_domains)
            if self.preferred_domains
            else "Not specified"
        )
        text = (
            f"- Title target: {self.title_target}\n"
            f"- Core skills: {skills}\n"
            f"- Soft skills (projects/education, limited professional experience): {soft}\n"
            f"- Years of experience: {self.years_experience}\n"
            f"- Location: {self.location}\n"
            f"- Preferred domains: {domains}\n"
            f"- Hard disqualifiers: {self.hard_disqualifiers}"
        )
        return Profile(text=text)


DISTILL_PROMPT = """You are distilling a candidate's resume into a short profile that a job-matching \
system will use to score job postings.

Extract the following from the resume text below:
- The seniority-appropriate title the candidate should target
- Their core skills: strengths backed by substantial professional (on-the-job) experience
- Their soft skills: skills demonstrated only through personal/academic projects, \
coursework, certifications, or education, WITHOUT strong professional experience \
(do not repeat anything already counted as a core skill)
- Total years of professional experience (estimate from the work history)
- Their location and any work-arrangement preferences
- Domains/industries they have worked in or clearly target
- Any hard constraints stated on the resume that should disqualify a posting

Base every field only on what the resume supports. Do not invent preferences \
that are not present; use "Not specified" or "None specified" where the resume \
is silent.

RESUME:
{resume_text}"""


def extract_docx_text(path: str) -> str:
    """Extract readable text (paragraphs and table cells) from a .docx file."""
    document = _DocxDocument(path)
    parts: List[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def distill_profile(
    resume_text: str,
    client: anthropic.Anthropic | None = None,
) -> Profile:
    """Distill raw resume text into a Profile via the LLM.

    Requires ANTHROPIC_API_KEY in the environment (or a preconfigured client).
    """
    if not resume_text.strip():
        raise ValueError("Resume text is empty; nothing to distill.")

    client = client or anthropic.Anthropic()
    response = client.messages.parse(
        model=MODEL,
        max_tokens=2048,
        thinking={"type": "adaptive"},
        messages=[
            {
                "role": "user",
                "content": DISTILL_PROMPT.format(resume_text=resume_text.strip()),
            }
        ],
        output_format=DistilledProfile,
    )

    distilled = response.parsed_output
    if distilled is None:
        raise RuntimeError(
            f"Model did not return a parseable profile (stop_reason={response.stop_reason})"
        )
    return distilled.to_profile()


def profile_from_docx(
    path: str,
    client: anthropic.Anthropic | None = None,
) -> Profile:
    """Read a .docx resume and distill it into a Profile."""
    text = extract_docx_text(path)
    if not text.strip():
        raise ValueError(f"No extractable text found in {path!r}.")
    return distill_profile(text, client=client)


def _main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("resume", help="Path to a .docx resume file.")
    parser.add_argument(
        "--set-default",
        action="store_true",
        help="Save the distilled profile as the new default (writes profile.txt), "
        "so future scorer runs use this resume.",
    )
    args = parser.parse_args()

    profile = profile_from_docx(args.resume)
    print(profile.text)

    if args.set_default:
        save_default_profile(profile.text)
        print(f"\nSaved as the new default profile: {PROFILE_PATH}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(_main())
