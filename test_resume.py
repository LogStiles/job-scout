"""Unit tests for resume.py (distiller rendering, docx extraction, errors).

The live distillation test makes a real LLM call and is skipped without a key.
"""

from __future__ import annotations

import os

import pytest
from docx import Document

import resume
import scorer

requires_api = pytest.mark.skipif(
    not os.environ.get("ANTHROPIC_API_KEY"),
    reason="ANTHROPIC_API_KEY not set; skipping live distillation test",
)


def _write_docx(path, paragraphs=(), table_rows=()):
    """Build a .docx at `path` from paragraphs and optional table rows."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_rows:
        cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=0, cols=cols)
        for row in table_rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = value
    doc.save(str(path))
    return str(path)


# --- DistilledProfile.to_profile -------------------------------------------


def _full_distilled(**overrides):
    base = dict(
        title_target="Senior Software Engineer",
        core_skills=["Java", "Spring Boot"],
        soft_skills=["Rust", "TensorFlow"],
        years_experience=4,
        location="NYC",
        preferred_domains=["fintech"],
        hard_disqualifiers="None specified",
    )
    base.update(overrides)
    return resume.DistilledProfile(**base)


def test_to_profile_returns_scorer_profile():
    profile = _full_distilled().to_profile()
    assert isinstance(profile, scorer.Profile)


def test_to_profile_renders_all_fields():
    text = _full_distilled().to_profile().text
    assert "- Title target: Senior Software Engineer" in text
    assert "- Core skills: Java, Spring Boot" in text
    assert "Rust, TensorFlow" in text
    assert "- Years of experience: 4" in text
    assert "- Location: NYC" in text
    assert "- Preferred domains: fintech" in text
    assert "- Hard disqualifiers: None specified" in text


def test_to_profile_soft_skills_line_labeled():
    text = _full_distilled().to_profile().text
    soft_line = next(l for l in text.splitlines() if l.startswith("- Soft skills"))
    assert "projects/education" in soft_line
    assert "Rust, TensorFlow" in soft_line


def test_to_profile_empty_soft_skills_is_not_specified():
    text = _full_distilled(soft_skills=[]).to_profile().text
    soft_line = next(l for l in text.splitlines() if l.startswith("- Soft skills"))
    assert soft_line.endswith("Not specified")


def test_to_profile_empty_core_and_domains_are_not_specified():
    text = _full_distilled(core_skills=[], preferred_domains=[]).to_profile().text
    assert "- Core skills: Not specified" in text
    assert "- Preferred domains: Not specified" in text


def test_to_profile_empty_string_fields_fall_back():
    # If the model returns blank strings, render sensible placeholders rather
    # than dangling "- Location: " lines.
    text = _full_distilled(
        title_target="   ", location="", hard_disqualifiers=""
    ).to_profile().text
    assert "- Title target: Not specified" in text
    assert "- Location: Not specified" in text
    assert "- Hard disqualifiers: None specified" in text


def test_to_profile_hard_disqualifiers_defaults():
    # hard_disqualifiers and soft_skills both have defaults.
    dp = resume.DistilledProfile(
        title_target="X", core_skills=["Go"], years_experience=2, location="Remote"
    )
    text = dp.to_profile().text
    assert "- Hard disqualifiers: None specified" in text
    assert "- Soft skills (projects/education, limited professional experience): Not specified" in text


def test_distilled_profile_rejects_negative_years():
    with pytest.raises(Exception):
        resume.DistilledProfile(
            title_target="X", core_skills=["Go"], years_experience=-1, location="NYC"
        )


# --- extract_docx_text ------------------------------------------------------


def test_extract_docx_text_reads_paragraphs(tmp_path):
    path = _write_docx(tmp_path / "r.docx", paragraphs=["First line", "Second line"])
    text = resume.extract_docx_text(path)
    assert "First line" in text
    assert "Second line" in text


def test_extract_docx_text_includes_table_cells(tmp_path):
    path = _write_docx(
        tmp_path / "r.docx",
        paragraphs=["Summary"],
        table_rows=[["Acme Corp", "Engineer, 2022-2025"]],
    )
    text = resume.extract_docx_text(path)
    assert "Summary" in text
    assert "Acme Corp | Engineer, 2022-2025" in text


def test_extract_docx_text_skips_empty_paragraphs(tmp_path):
    path = _write_docx(tmp_path / "r.docx", paragraphs=["Real", "   ", "", "Content"])
    text = resume.extract_docx_text(path)
    assert text == "Real\nContent"


def test_extract_docx_text_dedups_merged_cells(tmp_path):
    # A merged cell repeats across grid columns in row.cells; it should appear
    # once, not once per spanned column.
    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    merged = table.cell(0, 0).merge(table.cell(0, 2))
    merged.text = "Spanned Header"
    path = str(tmp_path / "merged.docx")
    doc.save(path)

    text = resume.extract_docx_text(path)
    assert text == "Spanned Header"


# --- error paths (no API needed) -------------------------------------------


def test_distill_profile_rejects_empty_text():
    with pytest.raises(ValueError):
        resume.distill_profile("   \n\t")


def test_profile_from_docx_rejects_empty_document(tmp_path):
    path = _write_docx(tmp_path / "empty.docx", paragraphs=[])
    with pytest.raises(ValueError):
        resume.profile_from_docx(path)


# --- live distillation ------------------------------------------------------


RESUME_TEXT = """Jane Doe — Backend Software Engineer — New York, NY (open to hybrid)

Professional experience:
- Acme Pay (2021-2025): 4 years building fintech payment microservices in
  Java and Spring Boot, deployed on Kubernetes and AWS.

Personal projects:
- Built a command-line note-taking tool in Rust (personal side project).

Education:
- B.S. Computer Science. Coursement included a machine-learning course using
  TensorFlow for a class project.
"""


@requires_api
def test_distill_profile_separates_core_and_soft_skills():
    profile = resume.distill_profile(RESUME_TEXT)
    assert isinstance(profile, scorer.Profile)
    text = profile.text
    lines = {
        "core": next(l for l in text.splitlines() if l.startswith("- Core skills:")),
        "soft": next(l for l in text.splitlines() if l.startswith("- Soft skills")),
    }
    # Java is clearly professional; Rust is a personal side project only.
    assert "Java" in lines["core"], text
    assert "Rust" in lines["soft"], text
    assert "Java" not in lines["soft"], text


@requires_api
def test_profile_from_docx_end_to_end(tmp_path):
    path = _write_docx(
        tmp_path / "jane.docx",
        paragraphs=RESUME_TEXT.splitlines(),
    )
    profile = resume.profile_from_docx(path)
    assert isinstance(profile, scorer.Profile)
    assert "- Title target:" in profile.text
    assert "- Soft skills" in profile.text
